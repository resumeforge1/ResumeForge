from __future__ import annotations

import re
import subprocess
import sys
import zipfile
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape
from playwright.sync_api import Error as PlaywrightError
from playwright.sync_api import sync_playwright

from app.template_registry import get_template_config


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "outputs"
TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
WINDOWS_BROWSER_PATHS = [
    Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
    Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
    Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
    Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
]

jinja_env = Environment(
    loader=FileSystemLoader(TEMPLATE_DIR),
    autoescape=select_autoescape(["html", "xml"]),
)


EXPORT_TYPES = {
    "ats_docx": ("ats-resume", "docx"),
    "premium_docx": ("premium-resume", "docx"),
    "premium_pdf": ("premium-resume", "pdf"),
    "cover_letter_docx": ("cover-letter", "docx"),
    "references_docx": ("references-sheet", "docx"),
    "linkedin_docx": ("linkedin-profile", "docx"),
    "interview_docx": ("interview-questions", "docx"),
    "application_email_txt": ("application-email", "txt"),
    "zip_package": ("complete-package", "zip"),
}


def slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return slug or "client"


def normalize_client(client: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(client)
    template_config = get_template_config(normalized.get("template_key"))
    normalized["template_config"] = template_config
    normalized["template_label"] = template_config["label"]
    normalized["experience_heading"] = template_config["experience_heading"]
    normalized["core_strengths"] = build_core_strengths(normalized)
    return normalized


def build_core_strengths(client: dict[str, Any]) -> list[str]:
    skills = client.get("skills") or []
    preferred = client.get("template_config", {}).get("core_strengths", [])
    strengths = [skill for skill in preferred if skill in skills]
    for skill in skills:
        if skill not in strengths:
            strengths.append(skill)
        if len(strengths) >= 9:
            break
    return strengths[:9]


def generate_outputs(client: dict[str, Any]) -> dict[str, str]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    client = normalize_client(client)

    return {
        export_type: generate_output(client, export_type)
        for export_type in EXPORT_TYPES
    }


def output_filename(client: dict[str, Any], export_type: str) -> str:
    label, extension = EXPORT_TYPES[export_type]
    return f"{int(client['id']):04d}-{slugify(client['full_name'])}-{label}.{extension}"


def generate_output(client: dict[str, Any], export_type: str) -> str:
    if export_type not in EXPORT_TYPES:
        raise ValueError(f"Unknown export type: {export_type}")
    OUTPUT_DIR.mkdir(exist_ok=True)
    client = normalize_client(client)
    filename = output_filename(client, export_type)
    path = OUTPUT_DIR / filename
    if export_type == "ats_docx":
        generate_ats_docx(client, path)
    elif export_type == "premium_docx":
        generate_premium_docx(client, path)
    elif export_type == "premium_pdf":
        generate_premium_pdf(client, path)
    elif export_type == "cover_letter_docx":
        generate_cover_letter_docx(client, path)
    elif export_type == "references_docx":
        generate_references_docx(client, path)
    elif export_type == "linkedin_docx":
        generate_linkedin_docx(client, path)
    elif export_type == "interview_docx":
        generate_interview_docx(client, path)
    elif export_type == "application_email_txt":
        generate_application_email(client, path)
    elif export_type == "zip_package":
        generate_zip_package(client, path)
    return filename


def generate_ats_docx(client: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(client["full_name"])
    run.bold = True
    run.font.size = Pt(18)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(" | ".join(filter(None, [client.get("city_state"), client.get("phone"), client.get("email")])))

    add_section(doc, "TARGET ROLE")
    doc.add_paragraph(client.get("target_role", ""))

    add_section(doc, "PROFESSIONAL SUMMARY")
    doc.add_paragraph(client.get("professional_summary", ""))

    add_bullet_section(doc, "CORE SKILLS", client.get("skills", []))
    add_bullet_section(doc, "LICENSES & CERTIFICATIONS", client.get("certifications", []))

    add_section(doc, client.get("experience_heading", "PROFESSIONAL EXPERIENCE").upper())
    for job in client.get("work_experience", []):
        heading = doc.add_paragraph()
        heading.add_run(job.get("job_title", "")).bold = True
        employer_line = " | ".join(
            filter(None, [job.get("employer"), f"{job.get('start_date', '')} - {job.get('end_date', '')}".strip(" -")])
        )
        if employer_line:
            heading.add_run(f" - {employer_line}")
        for bullet in job.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    add_bullet_section(doc, "EDUCATION", client.get("education", []))
    doc.save(path)


def add_section(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(21, 83, 130)


def add_bullet_section(doc: Document, title: str, items: list[str]) -> None:
    if not items:
        return
    add_section(doc, title)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def generate_cover_letter_docx(client: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    doc.add_paragraph(client["full_name"]).runs[0].bold = True
    doc.add_paragraph(" | ".join(filter(None, [client.get("city_state"), client.get("phone"), client.get("email")])))
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph("")
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph(
        f"I am excited to apply for the {client.get('target_role', 'open position')} opportunity. "
        f"My background aligns well with the needs of a team seeking dependable performance, clear communication, "
        f"and consistent attention to quality."
    )
    summary = client.get("professional_summary", "")
    if summary:
        doc.add_paragraph(summary)
    strengths = ", ".join(client.get("core_strengths", [])[:4])
    if strengths:
        doc.add_paragraph(
            f"I would bring practical strengths in {strengths}, along with a professional approach to daily responsibilities "
            f"and customer-facing work."
        )
    doc.add_paragraph(
        "Thank you for your time and consideration. I would welcome the opportunity to discuss how my experience can support your team."
    )
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(client["full_name"])
    doc.save(path)


def generate_premium_docx(client: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    heading = doc.add_paragraph()
    run = heading.add_run(client["full_name"])
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 87, 146)
    doc.add_paragraph(client.get("target_role", ""))
    doc.add_paragraph(" | ".join(filter(None, [client.get("city_state"), client.get("phone"), client.get("email")]))).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_section(doc, "PROFESSIONAL SUMMARY")
    doc.add_paragraph(client.get("professional_summary", ""))
    add_bullet_section(doc, "CORE STRENGTHS", client.get("core_strengths", []))
    add_section(doc, client.get("experience_heading", "PROFESSIONAL EXPERIENCE").upper())
    for job in client.get("work_experience", []):
        p = doc.add_paragraph()
        p.add_run(job.get("job_title", "")).bold = True
        p.add_run(f" | {job.get('employer', '')} | {job.get('start_date', '')} - {job.get('end_date', '')}")
        for bullet in job.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
    add_bullet_section(doc, "LICENSES & CERTIFICATIONS", client.get("certifications", []))
    add_bullet_section(doc, "EDUCATION", client.get("education", []))
    doc.save(path)


def generate_references_docx(client: dict[str, Any], path: Path) -> None:
    doc = Document()
    doc.add_heading(f"{client['full_name']} - References", level=1)
    doc.add_paragraph("References available upon request.")
    for index in range(1, 4):
        doc.add_heading(f"Reference {index}", level=2)
        doc.add_paragraph("Name:")
        doc.add_paragraph("Title / Relationship:")
        doc.add_paragraph("Company:")
        doc.add_paragraph("Phone:")
        doc.add_paragraph("Email:")
    doc.save(path)


def generate_linkedin_docx(client: dict[str, Any], path: Path) -> None:
    doc = Document()
    doc.add_heading(f"{client['full_name']} - LinkedIn Profile", level=1)
    doc.add_heading("Headline", level=2)
    doc.add_paragraph(client.get("target_role", ""))
    doc.add_heading("About", level=2)
    doc.add_paragraph(client.get("professional_summary", ""))
    doc.add_heading("Featured Skills", level=2)
    for skill in client.get("skills", [])[:20]:
        doc.add_paragraph(skill, style="List Bullet")
    doc.save(path)


def generate_interview_docx(client: dict[str, Any], path: Path) -> None:
    doc = Document()
    doc.add_heading(f"{client['full_name']} - Interview Questions", level=1)
    questions = [
        f"Tell me about your experience as a {client.get('target_role', 'professional')}.",
        "What safety, quality, or compliance standards do you follow in your work?",
        "Describe a time you solved a problem under pressure.",
        "How do you communicate with managers, customers, or teammates?",
        "Which strengths from your resume best match this role?",
    ]
    for question in questions:
        doc.add_paragraph(question, style="List Number")
    doc.save(path)


def generate_application_email(client: dict[str, Any], path: Path) -> None:
    body = (
        f"Subject: Application for {client.get('target_role', 'Open Position')}\n\n"
        "Dear Hiring Manager,\n\n"
        f"Please accept my application for the {client.get('target_role', 'open position')} role. "
        f"I have attached my resume and cover letter for your review.\n\n"
        "Thank you for your time and consideration.\n\n"
        f"Sincerely,\n{client['full_name']}\n"
    )
    path.write_text(body, encoding="utf-8")


def generate_zip_package(client: dict[str, Any], path: Path) -> None:
    included = [key for key in EXPORT_TYPES if key != "zip_package"]
    filenames = [generate_output(client, key) for key in included]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as package:
        for filename in filenames:
            package.write(OUTPUT_DIR / filename, arcname=filename)


def render_resume_html(client: dict[str, Any]) -> str:
    client = normalize_client(client)
    template_name = client["template_config"]["resume_template"]
    template = jinja_env.get_template(template_name)
    return template.render(client=client)


def generate_premium_pdf(client: dict[str, Any], path: Path) -> None:
    html = render_resume_html(client)
    render_pdf_with_playwright(html, path)


def render_pdf_with_playwright(html: str, path: Path) -> None:
    try:
        _render_pdf_with_browser(html, path)
    except PlaywrightError as exc:
        if "Executable doesn't exist" not in str(exc):
            raise
        browser_path = find_system_browser()
        if browser_path:
            _render_pdf_with_browser(html, path, browser_path)
            return
        subprocess.run(
            [sys.executable, "-m", "playwright", "install", "chromium"],
            check=True,
        )
        _render_pdf_with_browser(html, path)


def find_system_browser() -> Path | None:
    for browser_path in WINDOWS_BROWSER_PATHS:
        if browser_path.exists():
            return browser_path
    return None


def _render_pdf_with_browser(html: str, path: Path, executable_path: Path | None = None) -> None:
    with sync_playwright() as playwright:
        launch_options = {"executable_path": str(executable_path)} if executable_path else {}
        browser = playwright.chromium.launch(**launch_options)
        page = browser.new_page(viewport={"width": 816, "height": 1056})
        page.set_content(html, wait_until="networkidle")
        page.pdf(
            path=str(path),
            format="Letter",
            print_background=True,
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
        )
        browser.close()
