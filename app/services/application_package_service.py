from __future__ import annotations

import copy
import json
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

import app.document_generator as document_generator
from app.services.ai_providers import get_ai_provider
from app.services.dashboard_service import generate_interview_prep
from app.services.fresh_jobs import extract_candidate_profile, score_job, term_overlap
from app.services.job_analyzer import analyze_job_description


PACKAGE_EXPORT_TYPES = {
    "resume_docx": "resume-docx",
    "cover_letter_docx": "cover-letter-docx",
    "recruiter_email_txt": "recruiter-email-txt",
    "linkedin_message_txt": "linkedin-message-txt",
    "interview_questions_docx": "interview-questions-docx",
    "zip_package": "application-package-zip",
}


def build_application_package(client: dict[str, Any], job: dict[str, Any], existing_match: dict[str, Any] | None = None) -> dict[str, Any]:
    analysis = build_match_analysis(client, job, existing_match)
    tailored_resume = build_tailored_resume(client, analysis)
    interview = generate_interview_prep(client, job, {"matched_skills": analysis["strong_matching_skills"], "missing_qualifications": analysis["missing_keywords"]})
    role = job.get("title") or client.get("target_role") or "the role"
    company = job.get("company") or "the employer"
    return {
        "job_summary": {
            "company": company,
            "title": role,
            "location": job.get("location", ""),
            "description": job.get("description", ""),
        },
        "match_analysis": analysis,
        "resume_improvements": resume_improvements(client, analysis),
        "tailored_resume": tailored_resume,
        "tailored_resume_text": resume_draft_text(tailored_resume),
        "cover_letter": cover_letter_draft(client, job, analysis),
        "interview_questions": interview["questions"],
        "interview_summary": "\n".join(interview["talking_points"] + interview["star_prompts"]),
        "recruiter_email": recruiter_email_draft(client, job, analysis),
        "linkedin_message": linkedin_message_draft(client, job, analysis),
        "ats_keywords": analysis["ats_keywords"],
        "missing_skills_analysis": analysis["weak_areas"],
        "review_required": True,
        "guardrails": [
            "Review every draft before export.",
            "No application is submitted by ResumeForge.",
            "Employer names, dates, job titles, certificates, licenses, and numbers are preserved.",
        ],
    }


def build_match_analysis(client: dict[str, Any], job: dict[str, Any], existing_match: dict[str, Any] | None = None) -> dict[str, Any]:
    candidate_profile = extract_candidate_profile(client)
    match = existing_match or score_job(job, candidate_profile, {"target_role": client.get("target_role", ""), "location": client.get("city_state", "")})
    jd = analyze_job_description(job.get("description", ""), client)
    client_terms = candidate_profile.get("skills", []) + candidate_profile.get("certifications", [])
    missing_keywords = [
        keyword for keyword in jd.get("missing_keywords", [])
        if not any(term_overlap(keyword, item) for item in client_terms)
    ][:12]
    strong = match.get("matched_skills", []) or [
        skill for skill in candidate_profile.get("skills", [])
        if any(term_overlap(skill, keyword) for keyword in jd.get("keywords", []))
    ][:8]
    certs = jd.get("required_certifications", [])
    return {
        "match_score": int(match.get("score") or jd.get("ats_score") or 0),
        "match_percentage": int(jd.get("match_percentage") or match.get("qualification_coverage") or 0),
        "score_explanation": score_explanation(match),
        "breakdown": match.get("breakdown", {}),
        "missing_keywords": missing_keywords,
        "strong_matching_skills": strong[:8],
        "weak_areas": weak_areas(missing_keywords, certs, client),
        "certifications_mentioned": certs,
        "experience_alignment": experience_alignment(client, job),
        "education_alignment": education_alignment(client, job),
        "ats_keywords": dedupe((jd.get("keywords", []) + strong)[:24]),
    }


def score_explanation(match: dict[str, Any]) -> list[str]:
    breakdown = match.get("breakdown", {}) or {}
    if not breakdown:
        return ["Score is based on available job keywords and saved resume content."]
    return [f"{label.replace('_', ' ').title()}: {value}/100" for label, value in breakdown.items()]


def weak_areas(missing_keywords: list[str], certifications: list[str], client: dict[str, Any]) -> list[str]:
    output = [f"Keyword not clearly present: {keyword}" for keyword in missing_keywords[:6]]
    client_certs = " ".join(client.get("certifications", []) or []).lower()
    for cert in certifications:
        if cert.lower() not in client_certs:
            output.append(f"Certification requires user review before listing: {cert}")
    return output or ["No major weak areas identified from available data."]


def experience_alignment(client: dict[str, Any], job: dict[str, Any]) -> str:
    work = client.get("work_experience", []) or []
    if not work:
        return "No work history is saved, so experience alignment is unavailable."
    job_text = f"{job.get('title', '')} {job.get('description', '')}"
    titles = [entry.get("job_title", "") for entry in work]
    if any(term_overlap(title, job_text) for title in titles if title):
        return "Saved job titles and work history show partial or direct alignment with this posting."
    return "Work history exists, but direct title alignment is not clear from saved data."


def education_alignment(client: dict[str, Any], job: dict[str, Any]) -> str:
    education = client.get("education", []) or []
    job_text = job.get("description", "").lower()
    if not any(term in job_text for term in ("degree", "diploma", "education", "school", "ged")):
        return "The posting does not clearly state an education requirement."
    if education:
        return "Saved education is available for user review against the posting."
    return "The posting may mention education, but no education is saved."


def build_tailored_resume(client: dict[str, Any], analysis: dict[str, Any]) -> dict[str, Any]:
    provider = get_ai_provider()
    tailored = copy.deepcopy(client)
    original_summary = tailored.get("professional_summary", "")
    improved_summary = provider.rewrite(original_summary, "summary") if original_summary else ""
    tailored["professional_summary"] = improved_summary or original_summary
    existing_skills = tailored.get("skills", []) or []
    supported_keywords = [
        keyword for keyword in analysis.get("ats_keywords", [])
        if any(term_overlap(keyword, skill) for skill in existing_skills)
    ]
    tailored["skills"] = dedupe(existing_skills + supported_keywords)
    return tailored


def resume_improvements(client: dict[str, Any], analysis: dict[str, Any]) -> list[str]:
    improvements = []
    if analysis["missing_keywords"]:
        improvements.append("Review missing keywords and add only those already supported by real experience.")
    if analysis["strong_matching_skills"]:
        improvements.append("Emphasize strong matching skills in the summary and most relevant bullets.")
    if not client.get("professional_summary"):
        improvements.append("Add a truthful professional summary before exporting.")
    return improvements or ["Resume already has strong overlap with the posting; review wording before export."]


def cover_letter_draft(client: dict[str, Any], job: dict[str, Any], analysis: dict[str, Any]) -> str:
    provider = get_ai_provider()
    name = client.get("full_name", "Candidate")
    role = job.get("title") or client.get("target_role") or "the role"
    company = job.get("company") or "your team"
    skills = ", ".join(analysis.get("strong_matching_skills", [])[:4]) or "the qualifications listed in my resume"
    text = (
        f"{name}\n\nDear Hiring Manager,\n\n"
        f"I am excited to apply for the {role} opportunity with {company}. "
        f"My saved resume shows relevant strengths in {skills}. "
        "I have kept this draft review-only so every detail can be confirmed before submission.\n\n"
        "Thank you for your time and consideration.\n\n"
        f"Sincerely,\n{name}"
    )
    return provider.rewrite(text, "cover_letter")


def recruiter_email_draft(client: dict[str, Any], job: dict[str, Any], analysis: dict[str, Any]) -> str:
    role = job.get("title") or client.get("target_role") or "the role"
    company = job.get("company") or "your company"
    skills = ", ".join(analysis.get("strong_matching_skills", [])[:3]) or "the role requirements"
    return (
        f"Subject: Interest in {role}\n\n"
        f"Hello,\n\nI am interested in the {role} opportunity with {company}. "
        f"My resume includes experience aligned with {skills}. "
        "I would appreciate the opportunity to discuss whether my background fits your team's needs.\n\n"
        f"Thank you,\n{client.get('full_name', 'Candidate')}"
    )


def linkedin_message_draft(client: dict[str, Any], job: dict[str, Any], analysis: dict[str, Any]) -> str:
    role = job.get("title") or client.get("target_role") or "your opening"
    company = job.get("company") or "your team"
    skill = (analysis.get("strong_matching_skills", []) or ["my background"])[0]
    return (
        f"Hello, I saw the {role} opportunity with {company} and wanted to connect. "
        f"My resume includes relevant experience in {skill}. I would welcome a chance to learn more after reviewing the role."
    )


def resume_draft_text(client: dict[str, Any]) -> str:
    lines = [client.get("full_name", ""), client.get("target_role", ""), "", client.get("professional_summary", ""), "", "Skills:"]
    lines.extend(f"- {skill}" for skill in client.get("skills", []) or [])
    lines.append("\nExperience:")
    for entry in client.get("work_experience", []) or []:
        lines.append(f"{entry.get('job_title', '')} | {entry.get('employer', '')} | {entry.get('start_date', '')} - {entry.get('end_date', '')}")
        lines.extend(f"- {bullet}" for bullet in entry.get("bullets", []) or [])
    return "\n".join(line for line in lines if line is not None)


def package_from_form(form: Any, existing: dict[str, Any]) -> dict[str, Any]:
    package = copy.deepcopy(existing)
    for field in ("cover_letter", "recruiter_email", "linkedin_message", "interview_summary", "tailored_resume_text"):
        package[field] = str(form.get(field, package.get(field, "")))
    package["ats_keywords"] = split_lines(str(form.get("ats_keywords", "\n".join(package.get("ats_keywords", [])))))
    package["resume_improvements"] = split_lines(str(form.get("resume_improvements", "\n".join(package.get("resume_improvements", [])))))
    package["missing_skills_analysis"] = split_lines(str(form.get("missing_skills_analysis", "\n".join(package.get("missing_skills_analysis", [])))))
    package["interview_questions"] = split_lines(str(form.get("interview_questions", "\n".join(package.get("interview_questions", [])))))
    return package


def export_application_package(package_version: dict[str, Any], export_type: str) -> str:
    if export_type not in PACKAGE_EXPORT_TYPES:
        raise ValueError("Invalid package export type")
    document_generator.OUTPUT_DIR.mkdir(exist_ok=True)
    package = package_version["package"]
    client_id = int(package_version["client_id"])
    job_id = int(package_version["discovered_job_id"])
    base = f"{client_id:04d}-job-{job_id}-{PACKAGE_EXPORT_TYPES[export_type]}"
    if export_type == "resume_docx":
        filename = f"{base}.docx"
        document_generator.generate_ats_docx(package["tailored_resume"], document_generator.OUTPUT_DIR / filename)
    elif export_type == "cover_letter_docx":
        filename = f"{base}.docx"
        document_generator.generate_cover_letter_docx(
            package["tailored_resume"] | {"professional_summary": package["cover_letter"]},
            document_generator.OUTPUT_DIR / filename,
        )
    elif export_type == "interview_questions_docx":
        filename = f"{base}.docx"
        write_docx(document_generator.OUTPUT_DIR / filename, "Interview Questions", package.get("interview_questions", []))
    elif export_type == "recruiter_email_txt":
        filename = f"{base}.txt"
        (document_generator.OUTPUT_DIR / filename).write_text(package.get("recruiter_email", ""), encoding="utf-8")
    elif export_type == "linkedin_message_txt":
        filename = f"{base}.txt"
        (document_generator.OUTPUT_DIR / filename).write_text(package.get("linkedin_message", ""), encoding="utf-8")
    else:
        filename = f"{base}.zip"
        export_zip(package_version, document_generator.OUTPUT_DIR / filename)
    return filename


def export_zip(package_version: dict[str, Any], path: Path) -> None:
    exports = [key for key in PACKAGE_EXPORT_TYPES if key != "zip_package"]
    filenames = [export_application_package(package_version, export_type) for export_type in exports]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for filename in filenames:
            archive.write(document_generator.OUTPUT_DIR / filename, arcname=filename)


def write_docx(path: Path, title: str, lines: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(path)


def split_lines(value: str) -> list[str]:
    return [line.strip(" -\t") for line in value.splitlines() if line.strip(" -\t")]


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    output = []
    for item in items:
        key = item.strip().lower()
        if key and key not in seen:
            output.append(item)
            seen.add(key)
    return output
