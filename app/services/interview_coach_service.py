from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import app.document_generator as document_generator
from app.database import get_connection
from app.services.application_package_service import build_match_analysis
from app.services.dashboard_service import resume_readiness_score
from app.template_registry import get_template_config


INTERVIEW_MODES = [
    "General Interview",
    "Behavioral Interview",
    "Technical Interview",
    "Leadership Interview",
    "Manager Interview",
    "Customer Service",
    "Sales",
    "Warehouse",
    "Healthcare",
    "CDL Driver",
    "Software Engineer",
    "Finance",
    "Executive",
]


def build_interview_context(client: dict[str, Any] | None, job: dict[str, Any] | None, mode: str) -> dict[str, Any]:
    client = client or {}
    job = job or {}
    template = get_template_config(client.get("template_key")) if client else {"label": "General", "core_strengths": []}
    match = build_match_analysis(client, job) if client and job else {}
    return {
        "client": client,
        "job": job,
        "mode": mode if mode in INTERVIEW_MODES else "General Interview",
        "job_title": job.get("title") or client.get("target_role") or "the role",
        "industry": template.get("label", "General"),
        "skills": client.get("skills", []) or [],
        "certifications": client.get("certifications", []) or [],
        "match": match,
    }


def generate_questions(context: dict[str, Any]) -> list[dict[str, Any]]:
    mode = context["mode"]
    title = context["job_title"]
    skills = context.get("skills", [])[:8]
    base = [
        "Tell me about yourself.",
        f"Why do you want this {title} position?",
        "What makes you different from other candidates?",
        "How would you prioritize work when several tasks are urgent?",
        "Describe a difficult situation and how you handled it.",
        "Tell me about a time you received feedback.",
        "Describe a time you had to communicate clearly under pressure.",
        "What are your strongest skills for this role?",
        "What area are you actively improving?",
        "Why should we move you forward?",
    ]
    mode_questions = {
        "Behavioral Interview": [
            "Describe a conflict at work and how you resolved it.",
            "Tell me about a time you made a mistake and corrected it.",
            "Describe a time you had to adapt quickly.",
        ],
        "CDL Driver": [
            "Describe a safety issue you identified before it became a problem.",
            "How do you complete pre-trip and post-trip inspections?",
            "Describe your customer delivery experience.",
            "How do you stay compliant with DOT or company safety procedures?",
        ],
        "Customer Service": [
            "Describe a challenging customer interaction.",
            "How do you keep service professional during a delay or mistake?",
        ],
        "Leadership Interview": [
            "Describe a time you helped another person improve.",
            "How do you lead when you do not have formal authority?",
        ],
        "Manager Interview": [
            "How do you set expectations for a team?",
            "Describe a time you managed competing priorities.",
        ],
        "Sales": [
            "How do you build trust with a new prospect?",
            "Describe a time you handled an objection.",
        ],
        "Warehouse": [
            "How do you keep warehouse work accurate and safe?",
            "Describe your experience with inventory, loading, or equipment.",
        ],
        "Healthcare": [
            "How do you protect patient or client confidentiality?",
            "Describe a time you stayed calm during a high-pressure situation.",
        ],
        "Software Engineer": [
            "Describe a technical problem you debugged.",
            "How do you balance code quality and delivery speed?",
        ],
        "Finance": [
            "How do you verify accuracy in financial work?",
            "Describe a time you found and corrected an error.",
        ],
        "Executive": [
            "How do you define strategy and measure execution?",
            "Describe a time you led through uncertainty.",
        ],
        "Technical Interview": [
            "Describe the most technical part of your recent work.",
            "How do you learn a new tool or process?",
        ],
    }
    questions = base + mode_questions.get(mode, [])
    for skill in skills:
        questions.append(f"Tell me about your experience with {skill}.")
    while len(questions) < 15:
        questions.append(f"How does your background prepare you for {title}?")
    questions = dedupe(questions)[:30]
    return [{"text": question, "type": question_type(question), "star": star_guidance(question)} for question in questions]


def star_guidance(question: str) -> dict[str, str]:
    return {
        "Situation": "Set the real context briefly without adding new employers, dates, or claims.",
        "Task": "Explain your responsibility in that situation.",
        "Action": "Describe specific actions you personally took.",
        "Result": "Share the truthful outcome. If no metric exists, describe the practical result.",
    }


def review_answer(answer: str, question: str, context: dict[str, Any]) -> dict[str, Any]:
    words = re.findall(r"\b[\w'-]+\b", answer)
    lower = answer.lower()
    skills = [skill for skill in context.get("skills", []) if skill.lower() in lower]
    components = {
        "confidence": 80 if any(word in lower for word in ("i ", "my ", "led", "handled", "managed", "completed")) else 55,
        "specificity": 85 if any(char.isdigit() for char in answer) or len(skills) >= 2 else 60,
        "star_completeness": star_score(lower),
        "keywords": min(100, 45 + len(skills) * 12),
        "length": length_score(len(words)),
        "clarity": 85 if average_word_length(words) <= 7 else 68,
    }
    overall = round(sum(components.values()) / len(components))
    deductions = deductions_for(answer, words, components)
    return {
        **components,
        "overall": overall,
        "deductions": deductions,
        "strengths": strengths_for(components, skills),
        "weaknesses": deductions[:6],
        "suggested_improvements": improvements_for(deductions),
        "suggested_keywords": skills[:8],
        "missing_skills": [skill for skill in context.get("skills", [])[:8] if skill.lower() not in lower],
        "recommended_follow_up": "Revise the answer with a real STAR example and confirm every detail is accurate.",
    }


def start_session(client_id: int | None, job_id: int | None, mode: str, questions: list[dict[str, Any]]) -> int:
    with get_connection() as conn:
        cursor = conn.execute(
            """
            INSERT INTO interview_coach_sessions (client_id, discovered_job_id, mode, questions_json)
            VALUES (?, ?, ?, ?)
            """,
            (client_id, job_id, mode, json.dumps(questions)),
        )
        conn.commit()
        return int(cursor.lastrowid)


def get_session(session_id: int) -> dict[str, Any] | None:
    with get_connection() as conn:
        row = conn.execute("SELECT * FROM interview_coach_sessions WHERE id = ?", (session_id,)).fetchone()
    if not row:
        return None
    session = dict(row)
    session["questions"] = json.loads(session.get("questions_json") or "[]")
    session["answers"] = session_answers(session_id)
    session["summary"] = session_summary(session)
    return session


def session_answers(session_id: int) -> list[dict[str, Any]]:
    with get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM interview_coach_answers WHERE session_id = ? ORDER BY question_index",
            (session_id,),
        ).fetchall()
    answers = []
    for row in rows:
        item = dict(row)
        item["review"] = json.loads(item.get("review_json") or "{}")
        answers.append(item)
    return answers


def save_answer(session_id: int, question_index: int, answer: str, review: dict[str, Any]) -> None:
    with get_connection() as conn:
        conn.execute(
            """
            INSERT INTO interview_coach_answers (session_id, question_index, answer, review_json, completed)
            VALUES (?, ?, ?, ?, 1)
            ON CONFLICT(session_id, question_index)
            DO UPDATE SET answer = excluded.answer, review_json = excluded.review_json,
                completed = 1, updated_at = CURRENT_TIMESTAMP
            """,
            (session_id, question_index, answer, json.dumps(review)),
        )
        conn.commit()


def move_session(session_id: int, direction: str) -> None:
    session = get_session(session_id)
    if not session:
        return
    current = int(session.get("current_index") or 0)
    max_index = max(0, len(session["questions"]) - 1)
    next_index = min(max_index, current + 1) if direction == "next" else max(0, current - 1)
    with get_connection() as conn:
        conn.execute("UPDATE interview_coach_sessions SET current_index = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (next_index, session_id))
        conn.commit()


def mark_completed(session_id: int) -> None:
    with get_connection() as conn:
        conn.execute("UPDATE interview_coach_sessions SET completed = 1, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (session_id,))
        conn.commit()


def session_summary(session: dict[str, Any]) -> dict[str, Any]:
    answers = session.get("answers", [])
    scores = [int(answer.get("review", {}).get("overall") or 0) for answer in answers]
    best = max(answers, key=lambda item: item.get("review", {}).get("overall", 0), default=None)
    weakest = min(answers, key=lambda item: item.get("review", {}).get("overall", 100), default=None)
    average = round(sum(scores) / len(scores)) if scores else 0
    completion = round((len(answers) / max(1, len(session.get("questions", [])))) * 100)
    return {
        "questions_completed": len(answers),
        "average_score": average,
        "best_answer": best,
        "weakest_answer": weakest,
        "estimated_interview_readiness": round(average * 0.7 + completion * 0.3),
        "completion": completion,
    }


def interview_readiness(client: dict[str, Any] | None, session: dict[str, Any] | None, application_quality: int = 70) -> dict[str, Any]:
    resume_score = resume_readiness_score(client or {}).get("score", 0)
    answer_quality = (session or {}).get("summary", {}).get("average_score", 0)
    completion = (session or {}).get("summary", {}).get("completion", 0)
    components = {
        "resume_readiness": resume_score,
        "application_quality": application_quality,
        "answer_quality": answer_quality,
        "practice_completion": completion,
    }
    score = round(resume_score * 0.30 + application_quality * 0.20 + answer_quality * 0.35 + completion * 0.15)
    return {
        "score": score,
        "components": components,
        "explanation": ["30% resume readiness", "20% application quality", "35% answer quality", "15% practice completion"],
    }


def cheat_sheets() -> dict[str, list[str]]:
    return {
        "Questions to ask recruiter": ["What are the next steps?", "What does success look like in this role?", "What is the expected timeline?"],
        "Questions about salary": ["What is the listed compensation range?", "Is overtime or bonus eligibility documented?", "When are pay details finalized?"],
        "Questions about benefits": ["When do benefits begin?", "What training support is provided?", "Are schedules or routes consistent?"],
        "Questions about culture": ["How does the team communicate?", "How is safety or quality reinforced?", "What traits help someone succeed here?"],
        "Questions about growth": ["What advancement paths exist?", "How is performance reviewed?", "What additional certifications are valued?"],
    }


def body_language_tips() -> list[str]:
    return [
        "Maintain natural eye contact without staring.",
        "Use a steady speaking speed and pause before important details.",
        "Sit or stand with open posture.",
        "Dress professionally for the role and setting.",
        "Listen fully before answering.",
        "Send a truthful follow-up note after the interview.",
    ]


def follow_up_templates(client: dict[str, Any] | None, job: dict[str, Any] | None) -> dict[str, str]:
    name = (client or {}).get("full_name", "Candidate")
    role = (job or {}).get("title") or (client or {}).get("target_role") or "the role"
    company = (job or {}).get("company") or "your team"
    return {
        "Thank-you email": f"Dear Hiring Team,\n\nThank you for speaking with me about {role} with {company}. I appreciate your time and the chance to learn more.\n\nSincerely,\n{name}",
        "Recruiter follow-up": f"Hello,\n\nI am following up on my interview for {role}. I remain interested and would appreciate any update you can share.\n\nThank you,\n{name}",
        "Second interview response": f"Hello,\n\nThank you for inviting me to continue the interview process for {role}. I look forward to the next conversation.\n\n{name}",
        "Offer acceptance": f"Hello,\n\nThank you for the offer for {role}. I am pleased to accept, pending review of the final written details.\n\n{name}",
        "Offer decline": f"Hello,\n\nThank you for the offer for {role}. After review, I need to respectfully decline. I appreciate your time and consideration.\n\n{name}",
    }


def coach_dashboard_widgets(client_id: int | None = None) -> dict[str, Any]:
    with get_connection() as conn:
        sessions = conn.execute(
            "SELECT * FROM interview_coach_sessions WHERE (? IS NULL OR client_id = ?) ORDER BY updated_at DESC LIMIT 5",
            (client_id, client_id),
        ).fetchall()
        avg_row = conn.execute(
            """
            SELECT AVG(json_extract(review_json, '$.overall')) AS average_score
            FROM interview_coach_answers a
            JOIN interview_coach_sessions s ON s.id = a.session_id
            WHERE (? IS NULL OR s.client_id = ?)
            """,
            (client_id, client_id),
        ).fetchone()
    return {
        "recent_practice_sessions": [dict(row) for row in sessions],
        "average_interview_score": round(avg_row["average_score"] or 0),
        "practice_streak": len(sessions),
    }


def export_questions(session: dict[str, Any]) -> str:
    document_generator.OUTPUT_DIR.mkdir(exist_ok=True)
    filename = f"interview-session-{session['id']}-questions.txt"
    path = document_generator.OUTPUT_DIR / filename
    path.write_text("\n".join(question["text"] for question in session["questions"]), encoding="utf-8")
    record_export(session["id"], "questions_txt", filename)
    return filename


def export_notes_docx(session: dict[str, Any]) -> str:
    document_generator.OUTPUT_DIR.mkdir(exist_ok=True)
    filename = f"interview-session-{session['id']}-notes.docx"
    doc = Document()
    doc.add_heading("Interview Notes", level=1)
    for answer in session.get("answers", []):
        question = session["questions"][int(answer["question_index"])]["text"]
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer.get("answer", ""))
        doc.add_paragraph(f"Score: {answer.get('review', {}).get('overall', 0)}")
    doc.save(document_generator.OUTPUT_DIR / filename)
    record_export(session["id"], "notes_docx", filename)
    return filename


def export_summary_pdf(session: dict[str, Any]) -> str:
    document_generator.OUTPUT_DIR.mkdir(exist_ok=True)
    filename = f"interview-session-{session['id']}-summary.pdf"
    path = document_generator.OUTPUT_DIR / filename
    pdf = canvas.Canvas(str(path), pagesize=letter)
    text = pdf.beginText(72, 740)
    text.textLine("Interview Session Summary")
    text.textLine(f"Mode: {session.get('mode')}")
    text.textLine(f"Questions completed: {session['summary']['questions_completed']}")
    text.textLine(f"Average score: {session['summary']['average_score']}")
    text.textLine(f"Estimated readiness: {session['summary']['estimated_interview_readiness']}")
    pdf.drawText(text)
    pdf.save()
    record_export(session["id"], "summary_pdf", filename)
    return filename


def record_export(session_id: int, export_type: str, filename: str) -> None:
    with get_connection() as conn:
        conn.execute(
            "INSERT INTO interview_coach_exports (session_id, export_type, filename) VALUES (?, ?, ?)",
            (session_id, export_type, filename),
        )
        conn.commit()


def latest_session(client_id: int | None = None) -> dict[str, Any] | None:
    with get_connection() as conn:
        row = conn.execute(
            "SELECT id FROM interview_coach_sessions WHERE (? IS NULL OR client_id = ?) ORDER BY updated_at DESC LIMIT 1",
            (client_id, client_id),
        ).fetchone()
    return get_session(int(row["id"])) if row else None


def question_type(question: str) -> str:
    lower = question.lower()
    if any(term in lower for term in ("describe", "tell me about a time", "difficult", "conflict")):
        return "behavioral"
    if any(term in lower for term in ("technical", "tool", "process", "debug")):
        return "technical"
    return "general"


def star_score(lower_answer: str) -> int:
    score = 0
    for term in ("situation", "task", "action", "result"):
        if term in lower_answer:
            score += 25
    return score or 55


def length_score(word_count: int) -> int:
    if word_count < 35:
        return 50
    if word_count > 220:
        return 65
    return 90


def average_word_length(words: list[str]) -> float:
    return sum(len(word) for word in words) / max(1, len(words))


def deductions_for(answer: str, words: list[str], components: dict[str, int]) -> list[str]:
    deductions = []
    lower = answer.lower()
    if not any(char.isdigit() for char in answer):
        deductions.append("No measurable results.")
    if not any(word in lower for word in ("led", "improved", "resolved", "completed", "managed", "delivered", "trained")):
        deductions.append("Weak action verbs.")
    if "result" not in lower and not any(word in lower for word in ("improved", "reduced", "completed", "delivered", "resolved")):
        deductions.append("No outcome.")
    if not any(char.isdigit() for char in answer):
        deductions.append("No metrics.")
    if len(words) < 35:
        deductions.append("Too short.")
    if len(words) > 220:
        deductions.append("Too long.")
    if re.search(r"\b(was|were|been|being)\b", lower):
        deductions.append("Passive wording.")
    if components["specificity"] < 75:
        deductions.append("Generic wording.")
    return deductions


def strengths_for(components: dict[str, int], skills: list[str]) -> list[str]:
    strengths = []
    if components["length"] >= 80:
        strengths.append("Answer length is appropriate.")
    if components["star_completeness"] >= 75:
        strengths.append("STAR structure is visible.")
    if skills:
        strengths.append("Uses resume-supported keywords.")
    return strengths or ["Answer has a usable foundation for revision."]


def improvements_for(deductions: list[str]) -> list[str]:
    mapping = {
        "No measurable results.": "Add a truthful number only if one is available.",
        "Weak action verbs.": "Use direct action verbs such as handled, completed, resolved, or trained when accurate.",
        "No outcome.": "End with the real result or lesson learned.",
        "Too short.": "Add context, action, and outcome.",
        "Too long.": "Trim background details and focus on actions/results.",
        "Passive wording.": "Rewrite passive phrases into active, first-person wording.",
        "Generic wording.": "Add a specific example from saved experience.",
    }
    return [mapping.get(item, item) for item in deductions] or ["Practice once more and keep the answer factual."]


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    output = []
    for item in items:
        if item not in seen:
            output.append(item)
            seen.add(item)
    return output
